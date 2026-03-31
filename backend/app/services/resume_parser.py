import os
import re
from typing import Optional
import logging
# Move heavy imports inside to speed up startup
pdfplumber = None
PyPDF2 = None
Document = None

logger = logging.getLogger(__name__)

class ResumeParser:
    """Service for parsing resume files (PDF and DOCX)"""
    
    SUPPORTED_EXTENSIONS = {'.pdf', '.docx', '.doc'}
    
    def __init__(self):
        self.text_cleaners = [
            (r'\s+', ' '),  # Multiple spaces to single
            (r'\n\s*\n', '\n\n'),  # Multiple newlines to double
        ]
    
    def parse_file(self, file_path: str) -> Optional[str]:
        """Parse a resume file and extract text content"""
        if not os.path.exists(file_path):
            logger.error(f"File not found: {file_path}")
            return None
        
        ext = os.path.splitext(file_path)[1].lower()
        
        if ext not in self.SUPPORTED_EXTENSIONS:
            logger.warning(f"Unsupported file type: {ext}")
            return None
        
        try:
            if ext == '.pdf':
                return self._parse_pdf(file_path)
            elif ext in {'.docx', '.doc'}:
                return self._parse_docx(file_path)
        except Exception as e:
            logger.error(f"Error parsing file {file_path}: {str(e)}")
            return None
    
    def _parse_pdf(self, file_path: str) -> Optional[str]:
        """Extract text from PDF file"""
        text_content = []
        
        # Try pdfplumber first (better for complex layouts)
        try:
            global pdfplumber
            if pdfplumber is None:
                import pdfplumber as p
                pdfplumber = p
            
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_content.append(page_text)
            
            if text_content:
                return self._clean_text('\n'.join(text_content))
        except Exception as e:
            logger.warning(f"pdfplumber failed, trying PyPDF2: {str(e)}")
        
        # Fallback to PyPDF2
        try:
            global PyPDF2
            if PyPDF2 is None:
                import PyPDF2 as p
                PyPDF2 = p
                
            with open(file_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                for page in reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_content.append(page_text)
            
            return self._clean_text('\n'.join(text_content)) if text_content else None
        except Exception as e:
            logger.error(f"PyPDF2 also failed: {str(e)}")
            return None
    
    def _parse_docx(self, file_path: str) -> Optional[str]:
        """Extract text from DOCX file"""
        try:
            global Document
            if Document is None:
                from docx import Document as D
                Document = D
                
            doc = Document(file_path)
            text_content = []
            
            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text_content.append(para.text)
            
            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = ' | '.join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        text_content.append(row_text)
            
            return self._clean_text('\n'.join(text_content)) if text_content else None
        except Exception as e:
            logger.error(f"Error parsing DOCX: {str(e)}")
            return None
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize extracted text"""
        if not text:
            return ""
        
        # Apply text cleaners
        for pattern, replacement in self.text_cleaners:
            text = re.sub(pattern, replacement, text)
        
        return text.strip()
    
    def _extract_name(self, text: str) -> Optional[str]:
        """Extract candidate name from resume (usually in first few lines)"""
        if not text:
            return None
        
        # Get first 5 lines
        lines = [line.strip() for line in text.split('\n')[:5] if line.strip()]
        
        if not lines:
            return None
        
        # First line is often the name if it's all caps or title case
        first_line = lines[0]
        
        # Check if first line looks like a name (2-4 words, mostly letters)
        words = first_line.split()
        if 2 <= len(words) <= 4:
            # Check if it's mostly alphabetic (allowing for spaces and hyphens)
            if re.match(r'^[A-Za-z\s\-\.]+$', first_line):
                return first_line.title()
        
        # Try second line if first didn't work
        if len(lines) > 1:
            second_line = lines[1]
            words = second_line.split()
            if 2 <= len(words) <= 4 and re.match(r'^[A-Za-z\s\-\.]+$', second_line):
                return second_line.title()
        
        return None
    
    def extract_basic_info(self, text: str) -> dict:
        """Extract basic information using regex patterns"""
        info = {
            'name': self._extract_name(text),
            'email': None,
            'phone': None,
            'linkedin': None,
            'github': None
        }
        
        # Email pattern
        email_pattern = r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}'
        email_match = re.search(email_pattern, text)
        if email_match:
            info['email'] = email_match.group()
        
        # Phone pattern (various formats)
        phone_pattern = r'(?:\+\d{1,3}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}'
        phone_match = re.search(phone_pattern, text)
        if phone_match:
            info['phone'] = phone_match.group()
        
        # LinkedIn pattern
        linkedin_pattern = r'linkedin\.com/in/[\w-]+'
        linkedin_match = re.search(linkedin_pattern, text, re.IGNORECASE)
        if linkedin_match:
            info['linkedin'] = linkedin_match.group()
        
        # GitHub pattern
        github_pattern = r'github\.com/[\w-]+'
        github_match = re.search(github_pattern, text, re.IGNORECASE)
        if github_match:
            info['github'] = github_match.group()
        
        return info


# Singleton instance
resume_parser = ResumeParser()
